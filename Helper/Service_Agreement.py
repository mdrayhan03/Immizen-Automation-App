from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os

class Service_Agreement :
    def __init__(self, arr, current_path) :
        self.doc = Document()
        self.set_default_font('Calibri')
        self.set_margin(1)
        self.data_arr = arr

        self.path = ""
        if os.path.exists("path_file.txt") :
            f = open("path_file.txt", "rt")
            p = f.readline()
            self.path = p

        self.current_path = current_path
        self.logo()
        self.head = self.doc.add_heading("SERVICE AGREEMENT")
        self.head.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        self.first_part()
        self.milestone()
        self.second_part()
        self.table_part()
        self.third_part()
        self.contact()
        if self.data_arr[14] == True:
            self.reference()
        self.save()

    def set_default_font(self,font_name):
        styles = self.doc.styles
        style = styles['Normal']
        font = style.font
        font.name = font_name

    def set_margin(self, n) :
        section = self.doc.sections[0]
        section.left_margin = Inches(n)
        section.right_margin = Inches(n)
        section.top_margin = Inches(n)
        section.bottom_margin = Inches(n)
    
    def today(self) :
        today = datetime.today().date()
        d = today.strftime("%dth day of %B %Y")
        return d
    
    def logo(self) :        
        section = self.doc.sections[0]
        page_width = (section.page_width - section.left_margin - section.right_margin)
        header = section.header
        self.table = header.add_table(rows=1, cols=2, width=page_width)
        self.table.cell(0,0).text = ""
        img_run = self.table.cell(0,0).paragraphs[0]
        img_run.add_run().add_picture(f"assets/logo.png", width=Inches(2))
        self.table.cell(0,1).text = ""
        run = self.table.cell(0,1).paragraphs[0]
        run.add_run("Mailing Address: 502, 55 Commerce Valley Dr W\nMarkham, ON, L3T 7V9, Canada").font.size = Pt(8)
        run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self.table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    def first_part(self) :
        self.paragraph = self.doc.add_paragraph()
        self.run = self.paragraph.add_run(f"Client File Number: {self.data_arr[0]}")
        self.run.underline = True
        self.run.italic = True
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph(f"This Service Agreement is made this {self.today()}")
        self.doc.add_paragraph("BETWEEN\nSajid Iqbal (the “Regulated Canadian Immigration Consultant” or “RCIC”), RCIC License No: R712189located at Immizen Immigration Consulting Inc., 502- 55 Commerce Valley Dr W, Thornhill, ON, L3T 7V9, Canada").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph(f"AND\n{self.data_arr[1]}, Date of Birth: {self.data_arr[4]}, Passport No: {self.data_arr[7]}, located at address {self.data_arr[8]}").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("WHEREAS the RCIC and the Client wish to enter into a written agreement which contains the agreed-upon terms and conditions upon which the RCIC will provide his services to the Client.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("AND WHEREAS the RCIC is a licensee of the College of Immigration and Citizenship Consultants (“the College”), the regulator in Canada for immigration consultants; and RCIC is an Authorized Representative within the meaning of the Immigration and Refugee Protection Act (Canada) and the Citizenship Act (Canada) and the respective Government Regulations;").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("IN CONSIDERATION of the mutual covenants contained in this Agreement, the parties agree as follows:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        

    # def defination(self) :
        self.ol = self.doc.add_paragraph("Defination", "ListNumber")        
        self.paragraph = self.doc.add_paragraph()
        self.run = self.paragraph.add_run("The terms set out in this Service Agreement, have the meaning given to such terms in the By-law, Code of Professional Ethics, and Regulations of the College, as amended from time to time.")
        self.run.italic = True
        self.ol = self.doc.add_paragraph("RCIC Responsibilities and Commitments", "ListNumber")
        self.doc.add_paragraph(f"Both the RCIC and the Client agreed to apply for application {self.data_arr[9]}.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("RCIC will not bear any responsibility for any delay in submitting documents on time that might impact the status of the application of the applicant.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("In consideration of the fees paid and the matter stated above, the RCIC agrees to do the following:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY        
    
    def milestone(self) :
        self.doc.add_paragraph("The RCIC shall provide the Client with a finalized, signed copy of this Service Agreement.")
        i = 1
        for ele in self.data_arr[13] :
            self.doc.add_paragraph(f"Milestone-{i}", "ListBullet2")            
            self.doc.add_paragraph(ele[0])
            i += 1
    
    def second_part(self) :
    # def client_commitment(self) :
        self.doc.add_paragraph("Client Responsibilities and Commitments", "ListNumber")
        self.doc.add_paragraph("The Client must provide, upon request from the RCIC:", "ListNumber3")
        for ele in self.data_arr[12] :
            self.doc.add_paragraph(ele, "ListBullet2").paragraph_format.left_indent = Inches(1)

    # def third_part(self) :
        self.doc.add_paragraph("The Client understands that he/she must be accurate and honest in the information he/she provides and that any misrepresentations or omissions may void this Agreement, or seriously affect the outcome of the application or the retention of any immigration status he/she may obtain. The RCIC’s obligations under the Service Agreement are null and void if the Client knowingly provides any inaccurate, misleading, or false material information.  The Client’s financial obligations remain.", "ListNumber3").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("In the event Immigration, Refugees and Citizenship Canada (IRCC) or Employment and Social Development Canada (ESDC) or Provincial Government Administrator or processing Visa Office should contact the Client directly, the Client is instructed to notify the RCIC immediately.", "ListNumber3").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Client is to immediately advise the RCIC of any change in the marital, family, or civil status or change of physical address or contact information for any person included in the application.", "ListNumber3").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("In the event of a Joint Service Agreement, the Clients agree that the RCIC must share information among all clients, as required. Furthermore, if a conflict develops that cannot be resolved, the RCIC cannot continue to act for both or all of the Clients and may have to withdraw completely from representation.", "ListNumber3").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # def billing(self) :
        self.doc.add_paragraph("Billing Method", "ListNumber")
        self.doc.add_paragraph(f"The Client will be billed a total CAD ${self.data_arr[10]} (subject to taxes if applicable) in Canadian Currency to act as per the Service milestones or pre-determined schedule dates. Government and other charges & fees will be borne by the client. Additional CAD ${self.data_arr[11]}/hourly will be charged for any additional services upon mutual understanding.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # def payment(self) :
        self.doc.add_paragraph("Payment Term & Conditions", "ListNumber")
        self.doc.add_paragraph(f"Professional Fees	:    	CAD ${self.data_arr[10]}        - subject to relevant taxes, if any\nAdministrative Fees	: 	CAD ${self.data_arr[11]}         - applicable if cancelled")
        self.doc.add_paragraph("All payments are due upon receipt of the invoice and the first payment is on the execution of this Agreement. If the invoice amount is not received within 10 days of the invoice date, the agreement will be suspended, unless prior consent is obtained from the RCIC in writing; no further services will be provided until full payment is received. Failure to pay the invoice within 20 calendar days of the receipt date; the agreement will automatically terminate.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def table_part(self) :
    # def billing_schedule(self) :
        self.doc.add_paragraph("Billing Schedule", "ListNumber")
        self.table = self.doc.add_table(rows = len(self.data_arr[13])+2, cols = 4)
        self.table.style = "Table Grid"
        self.table.cell(0,0).text = "RCIC Service Milestone"
        self.table.cell(0,1).text = "Estimated date of completione"
        self.table.cell(0,2).text = "Professional Fees"
        self.table.cell(0,3).text = "Govt Fees/ Other Fees"
        
        tpro = 0
        tadm = 0
        for i in range(len(self.data_arr[13])) :
            self.table.cell(i+1,0).text = f"Milestone-{i+1}*"            
            self.table.cell(i+1,1).text = self.data_arr[13][i][1]
            self.table.cell(i+1,2).text = self.data_arr[13][i][2]
            self.table.cell(i+1,3).text = self.data_arr[13][i][3]
            tpro += int(self.data_arr[13][i][2])
            tadm += int(self.data_arr[13][i][3])

        i = len(self.data_arr[13])+1       
        self.table.cell(i,0).text = "Total"
        self.table.cell(i,1).text = ""
        self.table.cell(i,2).text = f"CAD ${tpro}"
        self.table.cell(i,3).text = f"CAD ${tadm}"

        self.paragraph = self.doc.add_paragraph()
        self.run = self.paragraph.add_run("*Milestones are defined on Section-2 of this agreement\n")
        self.run.italic = True
        self.run.font.size = Pt(9)

        self.run = self.paragraph.add_run(f"Client agrees to pay the full amount of CAD ${tpro}; professional fees and application fees in advance upon signing the agreement.").bold = True
        self.doc.add_paragraph(f"In any circumstances if the agreement is cancelled by the client CAD ${self.data_arr[11]} will be charged as Administrative Fees and any co-counseling fees paid to other representatives are not refundable.").alignment =WD_ALIGN_PARAGRAPH.JUSTIFY

    # def payment_method(self) :
        self.doc.add_paragraph("Payment Method", "ListNumber")
        self.doc.add_paragraph("Payment of professional fees/ other charges could be made by wire transfer, money order, interac/e- mail transfer, Stripe, Paypal, credit cards, direct deposit or cheques;").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Cheques and money order shall be made out to Immizen Immigration Consulting Inc.", "ListBullet2")
        self.doc.add_paragraph("Interac/email transfers should be made to payments@immizen.ca ", "ListBullet2")
        self.doc.add_paragraph("For Stripe and Paypal payments 3.6% additional charges will be applied ", "ListBullet2")
        self.doc.add_paragraph("CAD $17.5 will be charged extra for each Wire transfers/ international money transfer ", "ListBullet2")
        self.doc.add_paragraph("Payment by BDT is also accepted, whereas the exchange rate will be CAD $1= 90 BDT", "ListBullet2")
        self.doc.add_paragraph("Direct deposit/ Wire transfers shall be made to the following account:", "ListBullet2")
        self.doc.add_paragraph("Account Name: Immizen Immigration Consulting Inc.\nAccount No: 1009935\nTransit No: 02874\nInstitution No: 003\nBank: The Royal Bank of Canada (RBC)\nSWIFT Code: ROYCCAT2").paragraph_format.left_indent = Inches(0.5)
    
    def third_part(self) :
        self.doc.add_paragraph("Invoicing", "ListNumber")
        self.doc.add_paragraph("Invoices must be provided to the Client in accordance with the payment terms and conditions, found in section 5 of this Service Agreement. Additionally, upon the RCIC withdrawing or being discharged from representation, the RCIC must provide the Client with Statement of Account detailing all services that have been rendered or accounting for the time that has been spent on the Client’s file. ").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Refund Policy ", "ListNumber")
        self.doc.add_paragraph("The Client acknowledges that the granting of a visa or status and the time required for processing this application is at the sole discretion of the government of Canada (or Government Authorities) and not the RCIC.  Furthermore, the Client acknowledges that fees are not refundable in the event of an application refusal.  ").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY        
        self.doc.add_paragraph("If, however, the RCIC or professional staff do not complete the tasks identified under section 2 of this Agreement, the RCIC will refund part or all of the professional fees collected.  The Client agrees that the professional fees paid are for services indicated above, and any refund is strictly limited to the amount of professional fees paid.  Unused and/or unearned fees will be refunded in accordance with the Client File Management Regulation, the Client Account Regulation and the Service Agreement Regulation within 15 Business Day direct deposit to the client’s bank account or by cheque. ").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Co-counseling", "ListNumber")
        self.doc.add_paragraph("As required, the RCIC shall have the authority to assign co-counselors to engage in collaboration on the case, encompassing the exchange of client files, documents, and information exclusively for legal research purposes, and ensuring the preservation of legal practice standards to safeguard the client's confidentiality and privacy in accordance with the prevailing legal framework. The co-counselor, in compliance with this obligation, is expressly prohibited from disclosing the client's confidential information to any external party without obtaining the prior written consent of the designated representative or the client. The client hereby provides their consent to the RCIC for the assignment of co-counselors as and when required.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Dispute Resolution Related to the Code of Professional Ethics", "ListNumber")
        self.doc.add_paragraph("In the event of a dispute related to the Professional Services provided by the RCIC, the Client and RCIC are to make every reasonable effort to resolve the matter between the two parties.  In the event a resolution cannot be reached, the Client is to present the complaint in writing to the RCIC and allow the RCIC 30 days to respond to the Client.  In the event the dispute is still unresolved, the Client may follow the complaint and discipline procedure outlined by the College on their website: www.college-ic.ca").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The College Contact Information:\nCollege of Immigration and Citizenship\nConsultants (CICC)\n5500 North Service Rd., Suite 1002\nBurlington, ON, L7L 6W6\nToll-free: 1-877-836-7543").paragraph_format.left_indent = Inches(0.5)

        self.doc.add_paragraph("Confidentiality", "ListNumber")
        self.doc.add_paragraph("All information and documentation reviewed by the RCIC, required by IRCC and all other governing bodies, and used for the preparation of the application will not be divulged to any third party, other than agents and employees of the RCIC, without prior consent, except as demanded by the College or required under law.  The RCIC, and all agents and employees of the RCIC, are also bound by the confidentiality requirements of Article 8 of the Code of Professional Ethics.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Client agrees to the use of electronic communication and storage of confidential information.  The RCIC will use his/her best efforts to maintain a high degree of security for electronic communication and information storage.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Force Majeure", "ListNumber")
        self.doc.add_paragraph("The RCIC’s failure to perform any term of this Service Agreement, as a result of conditions beyond his/her control such as, but not limited to, governmental restrictions or subsequent legislation, war, strikes, or acts of God, shall not be deemed a breach of this Agreement.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Unplanned RCIC Absence", "ListNumber")
        self.doc.add_paragraph("In the event the Client is unable to contact the RCIC and has reason to believe the RCIC may be dead, incapacitated, or otherwise unable to fulfill his/her duties, the Client should contact the College.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Change Policy", "ListNumber")
        self.doc.add_paragraph("The Client acknowledges that if the RCIC is asked to act on the Client’s behalf on matters other than those outlined above in the scope of this Agreement, or because of a material change in the Client’s circumstances, or because of material facts not disclosed at the outset of the application, or because of a change in government legislation regarding the processing of immigration or citizenship-related applications, the Agreement can be modified accordingly.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("This Agreement may only be altered or amended when such changes are made in writing and executed by the parties hereto. All changes and/or edits must be initialled and dated by both the Licensee and the Client. Any substantial changes to this agreement may require that the parties enter into a new Service Agreement.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Termination", "ListNumber")
        self.doc.add_paragraph("This Agreement is considered terminated upon completion of tasks identified under section 2 of this agreement.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("This Agreement is considered terminated if material changes occur to the Client’s application or eligibility, which make it impossible to proceed with services detailed in section 2 of this Agreement.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Discharge or Withdrawal of Representation", "ListNumber")
        self.doc.add_paragraph("The Client may discharge representation and terminate this Agreement, upon writing, at which time any outstanding or unearned fees or Disbursements will be refunded by the RCIC to the Client and/or any outstanding fees or Disbursements will be paid by the Client to the RCIC.   ", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Pursuant to Article 11 of the Code of Professional Ethics, the RCIC may withdraw representation and terminate this Agreement, upon writing, provided withdrawal does not cause prejudice to the Client, at which time any outstanding or unearned fees or Disbursements will be refunded by the RCIC to the Client and/or any outstanding fees or Disbursements will be paid by the Client to the RCIC.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("At the time of withdrawal or discharge, the RCIC must provide the Client with an invoice detailing all services that have been rendered or accounting for the time that has been spent on the Client’s file.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Governing Law", "ListNumber")
        self.doc.add_paragraph("This Agreement shall be governed by the laws in effect in the Province of Ontario, and the federal laws of Canada applicable therein and except for disputes pursuant to Section 9 hereof, any dispute with respect to the terms of this Agreement shall be decided by a court of competent jurisdiction within the Province of Ontario.")

        self.doc.add_paragraph("Miscellaneous", "ListNumber")
        self.doc.add_paragraph("The Client expressly authorizes the RCIC to act on his behalf to the extent of the specific functions which the RCIC was retained to perform, as per Section 2 hereof.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("This Agreement constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes all prior agreements, understandings, warranties, representations, negotiations and discussions, whether oral or written, of the parties except as specifically set forth herein.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("This Agreement shall be binding upon the parties hereto and their respective heirs, administrators, successors and permitted assigns.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Costs enumerated in this Agreement are to be paid by the Client.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("This Agreement may only be altered or amended when such changes are made in writing and executed by the parties hereto. All changes and/or edits must be initialled and dated by both the Licensee and the Client. Any substantial changes to this Agreement may require that the parties enter into a new Service Agreement.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Client may, after a Service Agreement is signed, appoint a Designate to act on their behalf when dealing with the RCIC. A Designate must not be compensated by the Client or the RCIC for acting in the capacity of a Designate. ", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The provisions of this Agreement shall be deemed severable.  If any provision of this Agreement shall be held unenforceable by any court of competent jurisdiction, such provision shall be severed from this Agreement, and the remaining provisions shall remain in full force and effect.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The headings utilized in this Agreement are for convenience only and are not to be construed in any way as additions to or limitations of the covenants and agreements contained in this Agreement.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("Each of the parties hereto must do and execute or cause to be done or executed all such further and other things, acts, deeds, documents and assurances as may be necessary or reasonably required to carry out the intent and purpose of this Agreement fully and effectively.", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Client acknowledges that he has had sufficient time to review this Agreement and has been given an opportunity to obtain independent legal advice and translation prior to the execution and delivery of this Agreement. In the event the Client did not seek independent legal advice prior to signing this Agreement, he did so voluntarily without any undue pressure and agrees that the failure to obtain independent legal advice must not be used as a defence to the enforcement of obligations created by this Agreement.  ", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph("Furthermore, the Client acknowledges that he has received a copy of this Agreement and agrees to be bound by its terms.  ", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("The Client acknowledges that he has requested that the Agreement be written in the English language and that English is the binding language thereof;", "ListBullet2").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def contact(self) :
        self.doc.add_paragraph("Contact Information", "ListNumber")
        self.doc.add_paragraph(f"Client Information\nGiven Name: {self.data_arr[1]}\t\t\tFamily Name:  {self.data_arr[2]}\nAddress: {self.data_arr[8]}\nCellphone Number: {self.data_arr[5]}\nE-mail Address: {self.data_arr[6]}")

        self.doc.add_paragraph("RCIC Information\nGiven Name: Sajid\t\t\tFamily Name: Iqbal\nAddress: Immizen Immigration Consulting Inc., 502- 55 Commerce Valley Dr W Thornhill, ON, L3T 7V9, Canada\nCellphone Number: +1 4169490655\nE-mail Address: sajid@immizen.ca")

        self.doc.add_paragraph("IN WITNESS THEREOF this Agreement has been duly executed by the parties hereto on the signing date by both parties.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run("\n\n\n\n\n\n\n\n")
        self.paragraph.add_run("\t\t\t\t").underline = True
        self.paragraph.add_run("\t\t\t")
        self.paragraph.add_run("\t\t\t\t").underline = True
        self.paragraph.add_run("\n")
        self.paragraph.add_run('Client Name(the"Client")\t\t\t\tSajid Iqbal (the “RCIC”)')

    def reference(self) :
        self.doc.add_page_break()
        self.doc.add_paragraph(self.today())
        self.head = self.doc.add_heading(level=2)
        self.head.add_run("Reference: Assigning Designates").underline = True
        self.doc.add_paragraph(f"I, {self.data_arr[1]}, Date of Birth:  {self.data_arr[4]}, consent to designate {self.data_arr[15]} (email- {self.data_arr[16]}) to correspond, discuss, and share confidential information with Sajid Iqbal, RCIC on my behalf as required application {self.data_arr[9]} purpose.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run("\n\n\n\n\n\n\n\n")
        self.paragraph.add_run("\t\t\t\t\n").underline = True
        self.paragraph.add_run(f"{self.data_arr[1]}-Client\n")
        self.paragraph.add_run(f"Date of Birth: {self.data_arr[4]}\n")
        self.paragraph.add_run(f"Address: {self.data_arr[8]}")
        self.paragraph.add_run("\n\n\n\n\n\n\n\n")
        self.paragraph.add_run("\t\t\t\t\n").underline = True
        self.paragraph.add_run(f"{self.data_arr[15]}-Designate 1\n")
        self.paragraph.add_run(f"Date of Birth: {self.data_arr[17]}\n")
        self.paragraph.add_run(f"Address: {self.data_arr[18]}")
        
    def save(self) :
        self.doc.save(f"{self.path}/DocFile/Service Agreement/{self.data_arr[0]}.docx")

    def save_pdf(self) :
        import sys
        import logging
        sys.stdout = open('output.log', 'w')
        sys.stderr = open('error.log', 'w')
        logging.basicConfig(filename="app_error.log", level=logging.DEBUG)
        try:
            from docx2pdf import convert
            convert(f"{self.path}/DocFile/Service Agreement/{self.data_arr[0]}.docx")
        except Exception as e:
            logging.exception("An error occurred")
        