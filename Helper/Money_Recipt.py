from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
from num2words import num2words
import os

class Money_Receipt:
    def __init__(self, arr, current_path):
        self.doc = Document()
        self.set_default_font('Calibri')
        self.set_margin(1)
        self.data_arr = arr

        self.path = ""
        if os.path.exists("path_file.txt") :
            f = open("path_file.txt", "rt")
            p = f.readline()
            self.path = p

        self.current_path = current_path
        self.logo()
        self.header()
        self.table_part()
        self.save()

    def set_default_font(self, font_name):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name

    def set_margin(self, n):
        section = self.doc.sections[0]
        margin_in_inches = Inches(n)
        section.left_margin = margin_in_inches
        section.right_margin = margin_in_inches
        section.top_margin = margin_in_inches
        section.bottom_margin = margin_in_inches

    def today(self):
        return datetime.today().date().strftime("%B %d, %Y")

    def add_picture_to_cell(self, cell, image_path, width_in_inches):
        run = cell.paragraphs[0].add_run()
        run.add_picture(image_path, width=Inches(width_in_inches))

    def set_cell_text(self, cell, text, font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        paragraph.runs[0].font.size = font_size

    def logo(self):
        section = self.doc.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        header = section.header
        table = header.add_table(rows=1, cols=2, width=page_width)
        self.add_picture_to_cell(table.cell(0, 0), f"assets/logo.png", 2)
        self.set_cell_text(table.cell(0, 1), "Mailing Address: 502, 55 Commerce Valley Dr W\nMarkham, ON, L3T 7V9, Canada", Pt(8), WD_ALIGN_PARAGRAPH.RIGHT)
        table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def header(self):
        table = self.doc.add_table(rows=1, cols=2)
        self.set_cell_text(table.cell(0, 0), f"Date: {self.today()}")
        self.set_cell_text(table.cell(0, 1), f"Serial No: {self.data_arr[0]}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        self.doc.add_paragraph("\n")
        heading = self.doc.add_heading("MONEY RECEIPT\n\n")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 

    def table_part(self):
        table = self.doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"
        headers = ["Particulars", "Mode of Payment", "Total Amount (CAD)"]
        for i, header in enumerate(headers):
            self.set_cell_text(table.cell(0, i), header, alignment=WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)

        data = [
            (f"Client File No: {self.data_arr[1]}", "", ""),
            ("Consultation Fees", "", f"{self.data_arr[4]}"),
            (f"Application Fees", "", f"{self.data_arr[5]}"),
            ("", f"{self.data_arr[6]}", ""),
            ("Total:", "", f"{int(self.data_arr[4]) + int(self.data_arr[5])}")
        ]
        for i, (particulars, payment_mode, amount) in enumerate(data):
            self.set_cell_text(table.cell(i + 1, 0), particulars)
            self.set_cell_text(table.cell(i + 1, 1), payment_mode, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            self.set_cell_text(table.cell(i + 1, 2), amount, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        merge_cell = table.cell(1,1).merge(table.cell(2,1)).merge(table.cell(3,1))
        merge_cell.text = ""
        merge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        merge_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_picture_to_cell(merge_cell, f"assets/PAID.png", 1)

        self.doc.add_paragraph("\n\n\n")
        in_words = f"In Words: {num2words(int(self.data_arr[4]) + int(self.data_arr[5])).upper()} CAD only"
        self.doc.add_paragraph(in_words).bold = True
        self.doc.add_paragraph("\n\nPrepared By,")
        self.doc.add_paragraph("Immizen Immigration Consulting Inc.").bold = True

    def save(self):
        self.doc.save(f"{self.path}/DocFile/Money Receipt/{self.data_arr[0]}_{self.data_arr[1]}.docx")
    
    def save_pdf(self) :
        import sys
        import logging
        sys.stdout = open('output.log', 'w')
        sys.stderr = open('error.log', 'w')
        logging.basicConfig(filename="app_error.log", level=logging.DEBUG)
        try:
            from docx2pdf import convert
            convert(f"{self.path}/DocFile/Money Receipt/{self.data_arr[0]}_{self.data_arr[1]}.docx")
        except Exception as e:
            logging.exception("An error occurred")
        
