from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os

class Initial_Consultation :
    def __init__(self, arr, current_dir) :
        self.doc = Document()
        self.set_default_font('Calibri')
        self.set_margin(1)
        self.data_arr = arr
        self.current_dir = current_dir
        self.logo()
        self.head = self.doc.add_heading("INITIAL CONSULTATION AGREEMENT")
        self.head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.first_part()
        self.info()
        self.sign()
        self.footer()
        self.save()

    def set_default_font(self,font_name):
        styles = self.doc.styles
        style = styles['Normal']
        font = style.font
        font.name = font_name
    
    def set_margin(self, n) :
        section = self.doc.sections[0]
        section.left_margin = Inches(n)
        section.right_margin_margin = Inches(n)
        section.top_margin = Inches(n)
        section.bottom_margin = Inches(n)

    def today(self) :
        today = datetime.today().date()
        day = today.strftime("%B %d, %Y")
        return day

    def logo(self) :        
        section = self.doc.sections[0]
        page_width = (section.page_width - section.left_margin - section.right_margin)
        header = section.header
        self.table = header.add_table(rows=1, cols=2, width=page_width)
        self.table.cell(0,0).text = ""
        img_run = self.table.cell(0,0).paragraphs[0]
        img_run.add_run().add_picture(f"assets/logo.png", width=Inches(2))
        self.table.cell(0,1).text = ""
        run = self.table.cell(0,1).paragraphs[0]
        run.add_run("Mailing Address: 502, 55 Commerce Valley Dr W\nMarkham, ON, L3T 7V9, Canada").font.size = Pt(8)
        run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self.table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    def first_part(self) :
        self.doc.add_paragraph(f"This Initial Consultation Agreement is made this {self.today()}")
        self.doc.add_paragraph(f"between\nSajid Iqbal “Regulated Canadian Immigration Consultant (RCIC)”, License Number: R712189,\nand\n{self.data_arr[1]} the “Client”,").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph(f"for an Initial Consultation ({self.data_arr[10]}min) to discuss name the application {self.data_arr[9]}. The consultation is private and confidential. The client agrees to provide the required information during and before the assessment.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph(f"The client agrees to pay CAD {self.data_arr[11]} for the consultation, non-refundable.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run("This Agreement shall be governed by the laws in effect in the Province of Ontario, and the federal laws of Canada applicable therein.").italic = True
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run("Please be advised that Sajid Iqbal, RCIC is a member in good standing of the College of Immigration and Citizenship Consultants (CICC), and as such, is bound by its By-law, Code of Professional Ethics, and Regulations. ").italic = True
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def info(self) :
        self.heading = self.doc.add_heading(level=3)
        self.heading.add_run("Client's Information:\n").italic = True

        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run(f"Name		: {self.data_arr[1]}\nNationality	: {self.data_arr[3]}\nPassport No	: {self.data_arr[7]}\nDoB		: {self.data_arr[4]}\nEmail		: {self.data_arr[6]}\nPhone No	: {self.data_arr[5]}\nAddress   : {self.data_arr[8]}\n").italic = True
        
        self.heading = self.doc.add_heading(level=3)
        self.heading.add_run("RCIC’s Information:\n").italic = True

        self.paragraph = self.doc.add_paragraph()
        self.paragraph.add_run("Sajid Iqbal\nRCIC\nImmizen Immigration Consulting Inc.\nEmail: sajid@immizen.ca\n").italic = True

    def sign(self) :
        self.paragraph = self.doc.add_paragraph()
        self.run = self.paragraph.add_run("\n\t\t\t\t\t\t\t\tSajid Iqbal\n")
        self.run.italic = True
        self.run.font.size = Pt(22)
        self.run.font.name = "Edwardian Script ITC"

        self.paragraph.add_run("\t\t\t\t").underline = True
        self.paragraph.add_run("\t\t\t")
        self.paragraph.add_run("\t\t\t\t").underline = True

        self.doc.add_paragraph("Client's Signature\t\t\t\t\tRCIC’s  Signature")
    
    def footer(self) :
        section = self.doc.sections[0]
        footer = section.footer
        self.paragraph = footer.paragraphs[0]
        self.run = self.paragraph.add_run("Disclimer: This agreement covers only the 30-minute consultation and does not include client representation.")
        self.run.italic = True
        self.run.font.size = Pt(9)

    def save(self) :
        self.doc.save(f"DocFile/Initial Consultation Agreement/{self.data_arr[0]}.docx")
    
    def save_pdf(self) :
        import sys
        import logging
        sys.stdout = open('output.log', 'w')
        sys.stderr = open('error.log', 'w')
        logging.basicConfig(filename="app_error.log", level=logging.DEBUG)
        try:
            from docx2pdf import convert
            convert(f"DocFile/Initial Consultation Agreement/{self.data_arr[0]}.docx")
        except Exception as e:
            logging.exception("An error occurred")
        